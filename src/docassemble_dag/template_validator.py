"""
Template variable validation for DOCX, PDF, and Mako templates.

Validates that variables referenced in templates are defined in the
Docassemble interview, preventing production errors from undefined variables.
"""

import logging
import re
from pathlib import Path
from typing import Dict, List, Optional, Set, Tuple

from .graph import DependencyGraph

logger = logging.getLogger(__name__)

# Pattern to match Mako template variables ${variable} or ${object.attribute}
MAKO_VAR_PATTERN = re.compile(r'\$\{([a-zA-Z_][a-zA-Z0-9_]*(?:\.[a-zA-Z0-9_]+)?)\}')
# Pattern to match Docassemble template variables {variable} or {object.attribute}
# Supports nested attributes like {client.name.first}
TEMPLATE_VAR_PATTERN = re.compile(r'\{([a-zA-Z_][a-zA-Z0-9_]*(?:\.[a-zA-Z0-9_]+)*)\}')


class TemplateValidationResult:
    """Results of template variable validation."""
    
    def __init__(self, template_path: str) -> None:
        """Initialize validation result for a template file."""
        self.template_path = template_path
        self.extracted_variables: Set[str] = set()
        self.extracted_objects: Set[str] = set()
        self.undefined_variables: List[str] = []
        self.undefined_objects: List[str] = []
        self.valid_variables: List[str] = []
        self.valid_objects: List[str] = []
        self.is_valid = True
    
    def to_dict(self) -> Dict[str, any]:
        """Convert result to dictionary for JSON serialization."""
        return {
            "template_path": self.template_path,
            "extracted_variables": sorted(list(self.extracted_variables)),
            "extracted_objects": sorted(list(self.extracted_objects)),
            "undefined_variables": self.undefined_variables,
            "undefined_objects": self.undefined_objects,
            "valid_variables": self.valid_variables,
            "valid_objects": self.valid_objects,
            "is_valid": self.is_valid,
        }


def extract_template_variables(template_path: Path) -> Tuple[Set[str], Set[str]]:
    """
    Extract variable references from a template file.
    
    Supports DOCX, PDF, and Mako templates.
    
    Args:
        template_path: Path to template file
        
    Returns:
        Tuple of (set of variable names, set of object names)
        
    Raises:
        ValueError: If template format is not supported
        FileNotFoundError: If template file doesn't exist
    """
    if not template_path.exists():
        raise FileNotFoundError(f"Template file not found: {template_path}")
    
    suffix = template_path.suffix.lower()
    
    if suffix == '.docx':
        return _extract_docx_variables(template_path)
    elif suffix == '.pdf':
        return _extract_pdf_variables(template_path)
    elif suffix in ('.mako', '.html', '.txt'):
        return _extract_text_variables(template_path)
    else:
        raise ValueError(
            f"Unsupported template format: {suffix}. "
            "Supported formats: .docx, .pdf, .mako, .html, .txt"
        )


def _extract_docx_variables(template_path: Path) -> Tuple[Set[str], Set[str]]:
    """
    Extract variables from DOCX template.
    
    Attempts to parse DOCX and extract template variables.
    Falls back to text extraction if python-docx is not available.
    """
    variables: Set[str] = set()
    objects: Set[str] = set()
    
    try:
        # Try using python-docx if available
        try:
            from docx import Document
            
            doc = Document(template_path)
            text = '\n'.join(paragraph.text for paragraph in doc.paragraphs)
            # Also check table cells
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += '\n' + cell.text
        except ImportError:
            logger.warning(
                "python-docx not installed, extracting from DOCX as ZIP archive. "
                "Install with: pip install python-docx"
            )
            # Fallback: extract text from DOCX as ZIP
            import zipfile
            with zipfile.ZipFile(template_path, 'r') as zip_file:
                if 'word/document.xml' in zip_file.namelist():
                    xml_content = zip_file.read('word/document.xml').decode('utf-8', errors='ignore')
                    # Extract text between <w:t> tags
                    text = re.sub(r'<[^>]+>', ' ', xml_content)
                else:
                    text = ""
        
        variables, objects = _parse_template_text(text)
        
    except Exception as e:
        logger.error(f"Failed to extract variables from DOCX {template_path}: {e}")
        raise ValueError(f"Failed to parse DOCX template: {e}") from e
    
    return variables, objects


def _extract_pdf_variables(template_path: Path) -> Tuple[Set[str], Set[str]]:
    """
    Extract variables from PDF template.
    
    Attempts to extract text from PDF and find template variables.
    Falls back to basic text extraction if PDF libraries are not available.
    """
    variables: Set[str] = set()
    objects: Set[str] = set()
    
    try:
        # Try using pdfplumber if available
        try:
            import pdfplumber
            with pdfplumber.open(template_path) as pdf:
                text = '\n'.join(page.extract_text() or '' for page in pdf.pages)
        except ImportError:
            # Try PyPDF2 as fallback
            try:
                import PyPDF2
                with open(template_path, 'rb') as f:
                    pdf = PyPDF2.PdfReader(f)
                    text = '\n'.join(page.extract_text() or '' for page in pdf.pages)
            except ImportError:
                logger.warning(
                    "PDF parsing libraries not installed. "
                    "Install with: pip install pdfplumber or pip install PyPDF2"
                )
                text = ""
        
        variables, objects = _parse_template_text(text)
        
    except Exception as e:
        logger.error(f"Failed to extract variables from PDF {template_path}: {e}")
        raise ValueError(f"Failed to parse PDF template: {e}") from e
    
    return variables, objects


def _extract_text_variables(template_path: Path) -> Tuple[Set[str], Set[str]]:
    """Extract variables from plain text template (Mako, HTML, TXT)."""
    try:
        with open(template_path, 'r', encoding='utf-8') as f:
            text = f.read()
    except UnicodeDecodeError:
        # Try latin-1 as fallback
        with open(template_path, 'r', encoding='latin-1') as f:
            text = f.read()
    
    return _parse_template_text(text)


def _parse_template_text(text: str) -> Tuple[Set[str], Set[str]]:
    """
    Parse template text to extract variable references.
    
    Handles both Mako syntax ${variable} and Docassemble syntax {variable},
    as well as object attributes like {person.name} and nested {client.name.first}.
    """
    variables: Set[str] = set()
    objects: Set[str] = set()
    
    # Match Mako variables: ${variable} or ${object.attribute} or ${object.attribute.nested}
    for match in MAKO_VAR_PATTERN.findall(text):
        var_name = match
        if '.' in var_name:
            # Extract object name (first part before first dot)
            obj_name = var_name.split('.')[0]
            objects.add(obj_name)
            # Don't add intermediate parts as variables (e.g., "name" in "client.name.first")
        else:
            variables.add(var_name)
    
    # Match Docassemble template variables: {variable} or {object.attribute} or {object.attribute.nested}
    for match in TEMPLATE_VAR_PATTERN.findall(text):
        var_name = match
        if '.' in var_name:
            # Extract object name (first part before first dot)
            obj_name = var_name.split('.')[0]
            objects.add(obj_name)
            # Don't add intermediate parts as variables (e.g., "name" in "client.name.first")
        else:
            variables.add(var_name)
    
    return variables, objects


def validate_template(
    template_path: Path,
    interview_graph: DependencyGraph
) -> TemplateValidationResult:
    """
    Validate that all variables in a template are defined in the interview.
    
    Args:
        template_path: Path to template file
        interview_graph: Dependency graph from interview
        
    Returns:
        TemplateValidationResult with validation details
        
    Example:
        >>> graph = DependencyGraph(nodes, edges)
        >>> result = validate_template(Path("letter.docx"), graph)
        >>> if not result.is_valid:
        ...     print(f"Undefined variables: {result.undefined_variables}")
    """
    result = TemplateValidationResult(str(template_path))
    
    try:
        variables, objects = extract_template_variables(template_path)
        result.extracted_variables = variables
        result.extracted_objects = objects
        
        # Check if variables are defined in the interview
        interview_node_names = set(interview_graph.nodes.keys())
        
        for var_name in variables:
            if var_name in interview_node_names:
                result.valid_variables.append(var_name)
            else:
                result.undefined_variables.append(var_name)
                result.is_valid = False
        
        for obj_name in objects:
            if obj_name in interview_node_names:
                result.valid_objects.append(obj_name)
            else:
                result.undefined_objects.append(obj_name)
                result.is_valid = False
        
    except Exception as e:
        logger.error(f"Failed to validate template {template_path}: {e}")
        result.is_valid = False
        result.undefined_variables.append(f"ERROR: {e}")
    
    return result


def validate_templates(
    template_paths: List[Path],
    interview_graph: DependencyGraph
) -> Dict[str, TemplateValidationResult]:
    """
    Validate multiple templates against an interview graph.
    
    Args:
        template_paths: List of template file paths
        interview_graph: Dependency graph from interview
        
    Returns:
        Dictionary mapping template path to validation result
    """
    results: Dict[str, TemplateValidationResult] = {}
    
    for template_path in template_paths:
        try:
            result = validate_template(template_path, interview_graph)
            results[str(template_path)] = result
        except Exception as e:
            logger.warning(f"Failed to validate template {template_path}: {e}")
            result = TemplateValidationResult(str(template_path))
            result.is_valid = False
            result.undefined_variables.append(f"ERROR: {e}")
            results[str(template_path)] = result
    
    return results
